import streamlit as st
import pandas as pd
import re
from datetime import datetime, timedelta, date
import os
from typing import Dict, List, Optional, Any
from io import BytesIO, StringIO

class DocumentParser:
    """
    Handles parsing of various document formats including PDF, Word, Excel, and text.
    Extracts key information like deadlines, material requirements, and specifications.
    """
    
    def __init__(self):
        self.material_keywords = {
            'piping': ['pipe', 'piping', 'pipeline', 'tube', 'tubing'],
            'valves': ['valve', 'valves', 'ball valve', 'gate valve', 'check valve', 'control valve'],
            'flanges': ['flange', 'flanges', 'weld neck', 'slip on', 'blind flange'],
            'fittings': ['fitting', 'fittings', 'elbow', 'tee', 'reducer', 'coupling'],
            'bolts': ['bolt', 'bolts', 'stud', 'fastener', 'fasteners', 'screw'],
            'gaskets': ['gasket', 'gaskets', 'sealing', 'seal', 'o-ring'],
            'finned tubes': ['finned tube', 'finned tubes', 'fin tube', 'heat exchanger tube']
        }
        
        self.deadline_patterns = [
            r'deadline\s*:?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})',
            r'due\s*(?:by|on)?\s*:?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})',
            r'submit\s*(?:by|before)?\s*:?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})',
            r'no\s*later\s*than\s*:?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})',
            r'closing\s*date\s*:?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})',
            r'(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})\s*(?:is\s*the\s*)?deadline',
            r'(\w+\s+\d{1,2},?\s+\d{4})',  # December 15, 2024
            r'(\d{1,2}\s+\w+\s+\d{4})',    # 15 December 2024
        ]
        
        self.specification_keywords = [
            'specification', 'spec', 'requirement', 'standard', 'grade',
            'material', 'size', 'pressure', 'temperature', 'api', 'astm',
            'asme', 'din', 'en', 'iso', 'class', 'rating'
        ]
    
    def parse_file(self, uploaded_file) -> Dict[str, Any]:
        """
        Parse uploaded file and extract relevant information.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Dictionary containing extracted information
        """
        try:
            file_extension = uploaded_file.name.split('.')[-1].lower()
            
            if file_extension == 'pdf':
                return self._parse_pdf(uploaded_file)
            elif file_extension in ['docx', 'doc']:
                return self._parse_word(uploaded_file)
            elif file_extension in ['xlsx', 'xls']:
                return self._parse_excel(uploaded_file)
            elif file_extension == 'txt':
                return self._parse_text(uploaded_file.getvalue().decode('utf-8'))
            else:
                return {
                    'text': '',
                    'error': f'Unsupported file format: {file_extension}',
                    'materials': [],
                    'deadline': None,
                    'specifications': []
                }
        
        except Exception as e:
            return {
                'text': '',
                'error': f'Error parsing file: {str(e)}',
                'materials': [],
                'deadline': None,
                'specifications': []
            }
    
    def parse_text(self, text: str) -> Dict[str, Any]:
        """
        Parse plain text and extract information.
        
        Args:
            text: Plain text string
            
        Returns:
            Dictionary containing extracted information
        """
        return self._parse_text(text)
    
    def _parse_pdf(self, uploaded_file) -> Dict[str, Any]:
        """Parse PDF file using pdfplumber."""
        try:
            import pdfplumber
            
            text = ""
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if not text.strip():
                # If no text extracted, try OCR
                text = self._perform_ocr(uploaded_file)
            
            return self._analyze_text(text)
        
        except ImportError:
            return {
                'text': '',
                'error': 'pdfplumber not available. Cannot parse PDF files.',
                'materials': [],
                'deadline': None,
                'specifications': []
            }
        except Exception as e:
            return {
                'text': '',
                'error': f'Error parsing PDF: {str(e)}',
                'materials': [],
                'deadline': None,
                'specifications': []
            }
    
    def _parse_word(self, uploaded_file) -> Dict[str, Any]:
        """Parse Word document using python-docx."""
        try:
            from docx import Document
            
            doc = Document(uploaded_file)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return self._analyze_text(text)
        
        except ImportError:
            return {
                'text': '',
                'error': 'python-docx not available. Cannot parse Word documents.',
                'materials': [],
                'deadline': None,
                'specifications': []
            }
        except Exception as e:
            return {
                'text': '',
                'error': f'Error parsing Word document: {str(e)}',
                'materials': [],
                'deadline': None,
                'specifications': []
            }
    
    def _parse_excel(self, uploaded_file) -> Dict[str, Any]:
        """Parse Excel file using pandas."""
        try:
            # Try to read all sheets
            excel_file = pd.ExcelFile(uploaded_file)
            text = ""
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
                
                # Convert DataFrame to text
                text += f"Sheet: {sheet_name}\n"
                text += df.to_string(index=False) + "\n\n"
            
            return self._analyze_text(text)
        
        except Exception as e:
            return {
                'text': '',
                'error': f'Error parsing Excel file: {str(e)}',
                'materials': [],
                'deadline': None,
                'specifications': []
            }
    
    def _parse_text(self, text: str) -> Dict[str, Any]:
        """Parse plain text."""
        return self._analyze_text(text)
    
    def _perform_ocr(self, uploaded_file) -> str:
        """Perform OCR on uploaded file using pytesseract."""
        try:
            import pytesseract
            from PIL import Image
            import pdf2image
            
            # Convert PDF to images and perform OCR
            images = pdf2image.convert_from_bytes(uploaded_file.getvalue())
            text = ""
            
            for image in images:
                ocr_text = pytesseract.image_to_string(image)
                text += ocr_text + "\n"
            
            return text
        
        except ImportError:
            return "OCR libraries not available. Cannot extract text from scanned documents."
        except Exception as e:
            return f"OCR failed: {str(e)}"
    
    def _analyze_text(self, text: str) -> Dict[str, Any]:
        """
        Analyze extracted text to identify key information.
        
        Args:
            text: Extracted text
            
        Returns:
            Dictionary with analyzed information
        """
        text_lower = text.lower()
        
        # Extract materials
        materials = self._extract_materials(text_lower)
        
        # Extract deadline
        deadline = self._extract_deadline(text)
        
        # Extract specifications
        specifications = self._extract_specifications(text)
        
        # Extract project info
        project_info = self._extract_project_info(text)
        
        return {
            'text': text,
            'materials': materials,
            'deadline': deadline,
            'specifications': specifications,
            'project_name': project_info.get('project_name', ''),
            'tender_reference': project_info.get('tender_reference', '')
        }
    
    def _extract_materials(self, text_lower: str) -> List[str]:
        """Extract material categories from text."""
        found_materials = []
        
        for material, keywords in self.material_keywords.items():
            for keyword in keywords:
                if keyword in text_lower:
                    if material not in found_materials:
                        found_materials.append(material)
                    break
        
        return found_materials
    
    def _extract_deadline(self, text: str) -> Optional[datetime]:
        """Extract deadline from text using regex patterns."""
        for pattern in self.deadline_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                date_str = match.group(1)
                parsed_date = self._parse_date_string(date_str)
                if parsed_date:
                    return parsed_date
        
        return None
    
    def _parse_date_string(self, date_str: str) -> Optional[date]:
        """Parse various date string formats."""
        date_formats = [
            '%m/%d/%Y', '%d/%m/%Y', '%m-%d-%Y', '%d-%m-%Y',
            '%m.%d.%Y', '%d.%m.%Y', '%Y-%m-%d', '%Y/%m/%d',
            '%B %d, %Y', '%b %d, %Y', '%d %B %Y', '%d %b %Y'
        ]
        
        for fmt in date_formats:
            try:
                return datetime.strptime(date_str, fmt).date()
            except ValueError:
                continue
        
        return None
    
    def _extract_specifications(self, text: str) -> List[str]:
        """Extract technical specifications from text."""
        specifications = []
        lines = text.split('\n')
        
        for line in lines:
            line_lower = line.lower().strip()
            if not line_lower:
                continue
            
            # Check if line contains specification keywords or material-related terms
            is_specification = False
            for keyword in self.specification_keywords:
                if keyword in line_lower:
                    is_specification = True
                    break
            
            # Also check for lines with technical patterns (sizes, grades, standards)
            if not is_specification:
                technical_patterns = [
                    r'\d+["\']\s*(?:diameter|dia|pipe|tube)',  # Size specifications
                    r'(?:grade|class|schedule|rating)\s*[a-z0-9]+',  # Grades and standards
                    r'(?:api|ansi|astm|asme|iso)\s*[0-9a-z-]+',  # Standards
                    r'\d+\s*(?:mm|cm|in|inch|"|\')',  # Measurements
                    r'(?:carbon|stainless|alloy)\s*steel',  # Materials
                    r'(?:ball|gate|check|globe)\s*valve',  # Valve types
                ]
                for pattern in technical_patterns:
                    if re.search(pattern, line_lower):
                        is_specification = True
                        break
            
            if is_specification:
                # Clean and add the specification
                clean_spec = line.strip()
                if clean_spec and len(clean_spec) > 5:  # Reduced minimum length
                    specifications.append(clean_spec)
        
        # Remove duplicates while preserving order
        unique_specs = []
        for spec in specifications:
            if spec not in unique_specs:
                unique_specs.append(spec)
        
        return unique_specs  # Return all specifications found
    
    def _extract_project_info(self, text: str) -> Dict[str, str]:
        """Extract project name and tender reference from text."""
        project_info = {
            'project_name': '',
            'tender_reference': ''
        }
        
        lines = text.split('\n')
        
        # Patterns for project name
        project_patterns = [
            r'project\s*:?\s*(.+?)(?:\n|$)',
            r'project\s+name\s*:?\s*(.+?)(?:\n|$)',
            r'(?:title|name)\s*:?\s*(.+?)(?:\n|$)',
        ]
        
        # Patterns for tender reference
        reference_patterns = [
            r'(?:tender|ref|reference)\s*(?:no|number|#)?\s*:?\s*([A-Z0-9-]+)',
            r'(?:rfq|rfp|tender)\s*:?\s*([A-Z0-9-]+)',
            r'ref\s*:?\s*([A-Z0-9-]+)',
        ]
        
        text_lower = text.lower()
        
        # Extract project name
        for pattern in project_patterns:
            match = re.search(pattern, text_lower, re.IGNORECASE)
            if match:
                project_name = match.group(1).strip()
                if len(project_name) > 3 and len(project_name) < 100:
                    project_info['project_name'] = project_name.title()
                    break
        
        # Extract tender reference
        for pattern in reference_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                ref = match.group(1).strip()
                if len(ref) > 2:
                    project_info['tender_reference'] = ref.upper()
                    break
        
        return project_info
